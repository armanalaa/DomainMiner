from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


OUT = Path("La Dispensa.docx").resolve()
IMG_DIR = Path("ppt_extract_images").resolve()


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=120, start=160, bottom=120, end=160):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def add_blank(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    return p


def set_run_font(run, name, size=14, color=None, bold=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run._element.rPr.rFonts.set(qn("w:cs"), name)
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold


def add_heading(doc, text, level=1, numbered=True, prefix=""):
    add_blank(doc)
    if level == 1:
        style = "Heading 2"
    elif level == 2:
        style = "Heading 3"
    else:
        style = "Heading 4"
    p = doc.add_paragraph(style=style)
    run = p.add_run((prefix + " " if prefix else "") + text)
    if level == 1:
        set_run_font(run, "Aptos Display", 16, "0F4761", True)
    elif level == 2:
        set_run_font(run, "Aptos", 14, "0F4761", True)
    else:
        set_run_font(run, "Aptos", 12, "0F4761", True)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_blank(doc)
    return p


def add_exercise_heading(doc):
    add_blank(doc)
    p = doc.add_paragraph()
    r = p.add_run("Esercizi")
    set_run_font(r, "Aptos", 14, "000000", True)
    add_blank(doc)


def add_para(doc, parts):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.15
    if isinstance(parts, str):
        parts = [(parts, False)]
    for text, code in parts:
        r = p.add_run(text)
        if code:
            set_run_font(r, "Courier New", 14, "1F4E79", False)
        else:
            set_run_font(r, "Times New Roman", 14, "000000", False)
    return p


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0000FF")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), "Times New Roman")
    r_fonts.set(qn("w:hAnsi"), "Times New Roman")
    r_fonts.set(qn("w:eastAsia"), "Times New Roman")
    r_fonts.set(qn("w:cs"), "Times New Roman")
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), "24")
    r_pr.extend([r_fonts, color, underline, size])
    new_run.append(r_pr)
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    set_run_font(r, "Times New Roman", 14, "000000", False)
    return p


def add_code(doc, code):
    add_blank(doc)
    table = doc.add_table(rows=1, cols=1)
    table.autofit = True
    cell = table.cell(0, 0)
    set_cell_shading(cell, "EAF2F8")
    set_cell_margins(cell)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    for idx, line in enumerate(code.strip("\n").split("\n")):
        if idx:
            p.add_run().add_break()
        r = p.add_run(line)
        set_run_font(r, "Courier New", 14, "1F4E79", False)
    add_blank(doc)


def add_figure(doc, image_name, title, caption, width_cm=14.5):
    add_blank(doc)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run()
    r.add_picture(str(IMG_DIR / image_name), width=Cm(width_cm))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(8)
    run = cap.add_run(f"Figura - {caption}. Titolo della risorsa: {title}.")
    set_run_font(run, "Times New Roman", 12, "000000", False)
    add_blank(doc)


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:cs"), "Times New Roman")
    normal.font.size = Pt(14)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.15

    for name, font, size in [
        ("Heading 2", "Aptos Display", 16),
        ("Heading 3", "Aptos", 14),
        ("Heading 4", "Aptos", 12),
    ]:
        st = styles[name]
        st.font.name = font
        st._element.rPr.rFonts.set(qn("w:eastAsia"), font)
        st._element.rPr.rFonts.set(qn("w:cs"), font)
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor(0x0F, 0x47, 0x61)
        st.paragraph_format.space_before = Pt(0)
        st.paragraph_format.space_after = Pt(0)

    for name in ["List Bullet", "List Number"]:
        st = styles[name]
        st.font.name = "Times New Roman"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        st.font.size = Pt(14)
        st.paragraph_format.space_after = Pt(6)


def build():
    doc = Document()
    configure_document(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run("La Dispensa")
    set_run_font(tr, "Aptos Display", 18, "0F4761", True)
    add_blank(doc)
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run("Input utente, accesso ai dati e strutture nidificate nei dizionari")
    set_run_font(sr, "Times New Roman", 14, "000000", False)
    add_blank(doc)

    add_heading(doc, "Dati inseriti dall'utente e dizionari", 1, prefix="1.")
    add_para(doc, "Un programma diventa più utile quando non lavora soltanto con valori scritti direttamente nel codice, ma riceve informazioni dall'utente durante l'esecuzione. In questo contesto, il dizionario permette di raccogliere più dati collegati tra loro usando chiavi descrittive. Il caso dei dati anagrafici mostra bene questa idea, perché nome, cognome ed età appartengono alla stessa persona ma restano accessibili separatamente.")
    add_para(doc, [("Il punto di partenza è un dizionario vuoto, scritto come ", False), ("persona = {}", True), (", che viene riempito progressivamente. Ogni dato letto dalla tastiera viene assegnato a una chiave precisa, come ", False), ("\"nome\"", True), (", ", False), ("\"cognome\"", True), (" ed ", False), ("\"eta\"", True), (". In questo modo il programma non memorizza una sequenza anonima di valori, ma una piccola struttura in cui il significato di ogni informazione è esplicito.", False)])
    add_code(doc, '''persona = {}
print("Inserire i seguenti dati anagrafici:")
persona["nome"] = input("Nome: ")
persona["cognome"] = input("Cognome: ")
persona["eta"] = input("Età: ")
print("I dati inseriti sono:", persona)''')
    add_para(doc, [("La prima istruzione crea il dizionario vuoto, quindi prepara il contenitore nel quale saranno inseriti i dati. Le tre istruzioni con ", False), ("input()", True), (" chiedono all'utente di scrivere nome, cognome ed età, e ciascuna risposta viene salvata nella chiave corrispondente. L'ultima istruzione usa ", False), ("print()", True), (" per mostrare l'intero dizionario e consente di verificare che i dati siano stati memorizzati con le chiavi previste.", False)])
    add_figure(doc, "slide_04_pic_1.png", "Esempio: costruzione di un dizionario (2/3)", "Esecuzione in IDLE della costruzione di un dizionario", 13.5)

    add_heading(doc, "Il valore prodotto da input()", 2, prefix="I.")
    add_para(doc, [("Il passaggio successivo è osservare il tipo dei valori ottenuti con ", False), ("input()", True), (". Anche quando l'utente scrive un numero, il valore letto viene trattato come stringa, quindi l'età inserita come ", False), ("57", True), (" viene salvata come testo se non viene convertita. Questa distinzione è importante perché il dizionario può contenere il valore, ma il significato operativo del valore dipende anche dal suo tipo.", False)])
    add_para(doc, [("Per memorizzare l'età come numero intero, il materiale usa la conversione con ", False), ("int(input(...))", True), (". L'idea è semplice: prima viene letto il testo scritto dall'utente, poi quel testo viene trasformato in un intero. Il dizionario finale conserva la stessa chiave ", False), ("\"eta\"", True), (", ma il valore associato è ora un numero e non una stringa.", False)])
    add_code(doc, '''persona = {}
persona["eta"] = int(input("Età: "))''')
    add_para(doc, [("La prima riga crea ancora il dizionario vuoto, quindi non contiene dati iniziali. Nella seconda riga, ", False), ("input(\"Età: \")", True), (" acquisisce ciò che l'utente scrive dopo il messaggio mostrato nella shell. La funzione ", False), ("int()", True), (" converte quel valore in un numero intero prima che venga assegnato alla chiave ", False), ("\"eta\"", True), (".", False)])
    add_figure(doc, "slide_05_pic_1.png", "Esempio: costruzione di un dizionario (3/3)", "Esecuzione in IDLE con età convertita in numero intero", 13.5)

    add_heading(doc, "Accesso ai dati mediante chiave", 1, prefix="2.")
    add_para(doc, "Dopo avere costruito un dizionario, il problema naturale diventa leggere singole informazioni senza stampare l'intera struttura. L'accesso avviene usando il nome del dizionario seguito dalla chiave tra parentesi quadre. Questa forma rende chiaro che non si sta cercando un elemento per posizione, ma per significato.")
    add_para(doc, [("Se il dizionario contiene le chiavi ", False), ("\"nome\"", True), (", ", False), ("\"cognome\"", True), (" e ", False), ("\"eta\"", True), (", ciascun valore può essere recuperato separatamente. Il materiale presenta questa operazione dentro una funzione, perché la funzione riceve il dizionario come argomento e stampa un resoconto ordinato. In questo modo la struttura dei dati e il modo di presentarli restano collegati ma distinguibili.", False)])
    add_code(doc, '''def stampa_dizionario(persona):
    print("I dati anagrafici sono i seguenti:")
    print("Nome:", persona["nome"])
    print("Cognome:", persona["cognome"])
    print("Età:", persona["eta"])''')
    add_para(doc, [("La riga con ", False), ("def", True), (" definisce una funzione che si chiama ", False), ("stampa_dizionario", True), (" e che riceve un parametro chiamato ", False), ("persona", True), (". Le tre istruzioni successive accedono ai valori del dizionario usando le chiavi e li affiancano a etichette leggibili nella shell. La funzione non costruisce il dizionario, ma si limita a usare una struttura già esistente e quindi mostra come l'accesso per chiave renda i dati riutilizzabili.", False)])

    add_heading(doc, "Dati, funzione e messaggi chiari", 2, prefix="II.")
    add_para(doc, [("La funzione di stampa introduce anche una scelta didattica importante: il programma non deve soltanto produrre valori, ma deve comunicarli in modo comprensibile. Il messaggio iniziale spiega che saranno mostrati dati anagrafici, mentre ogni riga successiva associa un'etichetta al valore estratto. Questa organizzazione evita che l'utente debba interpretare da solo la struttura grezza del dizionario.", False)])
    add_para(doc, [("Il collegamento con il primo esempio è diretto, perché le stesse chiavi usate per memorizzare i dati sono usate per recuperarli. Se una chiave cambia durante la costruzione del dizionario, anche la funzione che la legge deve usare lo stesso nome. La coerenza tra inserimento e accesso è quindi parte essenziale della correttezza del programma.", False)])

    add_heading(doc, "Dizionari per rappresentare vettori", 1, prefix="3.")
    add_para(doc, "Una volta compreso che un dizionario può raccogliere dati collegati, la stessa struttura può rappresentare oggetti più astratti. Il materiale applica i dizionari ai vettori, descrivendo un vettore mediante due componenti. La componente orizzontale viene associata alla chiave x, mentre la componente verticale viene associata alla chiave y.")
    add_para(doc, [("La forma scelta è ", False), ("{'x': valore_x, 'y': valore_y}", True), (", cioè un dizionario con due coppie chiave-valore. Questa rappresentazione è compatta ma leggibile, perché le chiavi indicano il ruolo delle componenti. Due vettori possono quindi essere salvati come due dizionari distinti, ad esempio ", False), ("v1", True), (" e ", False), ("v2", True), (", ciascuno con le proprie componenti.", False)])
    add_code(doc, '''v1 = {'x': 2, 'y': 3}
v2 = {'x': -1, 'y': 5}''')
    add_para(doc, [("Il primo dizionario rappresenta un vettore con componente ", False), ("x", True), (" uguale a ", False), ("2", True), (" e componente ", False), ("y", True), (" uguale a ", False), ("3", True), (". Il secondo rappresenta un altro vettore, con componente orizzontale negativa e componente verticale positiva. La scelta delle chiavi permette di sommare le componenti corrispondenti senza confondere il ruolo dei valori.", False)])

    add_heading(doc, "Somma componente per componente", 2, prefix="I.")
    add_para(doc, "La somma vettoriale presentata nel materiale è un esempio concreto di elaborazione di dizionari. Il programma riceve due dizionari che rappresentano vettori, crea un nuovo dizionario e inserisce nel risultato la somma delle componenti omonime. Il risultato non modifica i vettori di partenza, ma produce una nuova struttura con le chiavi x e y.")
    add_code(doc, '''def somma_vettoriale(v1, v2):
    risultato = {}
    risultato['x'] = v1['x'] + v2['x']
    risultato['y'] = v1['y'] + v2['y']
    return risultato''')
    add_para(doc, [("La funzione riceve due parametri, ", False), ("v1", True), (" e ", False), ("v2", True), (", che devono avere la stessa struttura: una chiave ", False), ("'x'", True), (" e una chiave ", False), ("'y'", True), (". Il dizionario ", False), ("risultato", True), (" viene inizialmente creato vuoto, poi riempito con la somma delle componenti orizzontali e verticali. L'istruzione ", False), ("return", True), (" restituisce il nuovo dizionario, rendendo il valore calcolato disponibile al resto del programma.", False)])
    add_figure(doc, "slide_12_pic_1.jpg", "Esempio: somma vettoriale (4/2)", "Shell e funzione per la somma vettoriale", 14.5)

    add_heading(doc, "Strutture nidificate nei dizionari", 1, prefix="4.")
    add_para(doc, "I dizionari diventano più espressivi quando i valori associati alle chiavi non sono soltanto numeri o stringhe. Un valore può essere una lista, oppure può essere un altro dizionario. Questa possibilità introduce le strutture nidificate, cioè strutture di dati contenute dentro altre strutture di dati.")
    add_para(doc, [("Il collegamento con gli esempi precedenti è naturale: prima una chiave conteneva un singolo dato, ora una chiave può contenere un gruppo di dati. Se una chiave contiene una lista, si accede prima alla chiave e poi alla posizione nella lista. Se una chiave contiene un dizionario, si accede prima alla chiave esterna e poi alla chiave interna.", False)])
    add_code(doc, '''d = {"numeri": [10, 20, 30]}
print(d["numeri"][1])

d = {"persona": {"nome": "Anna", "età": 22}}
print(d["persona"]["nome"])''')
    add_para(doc, [("Nel primo esempio, la chiave ", False), ("\"numeri\"", True), (" contiene una lista con tre valori, quindi ", False), ("d[\"numeri\"]", True), (" recupera la lista e ", False), ("[1]", True), (" seleziona il secondo elemento, cioè ", False), ("20", True), (". Nel secondo esempio, la chiave ", False), ("\"persona\"", True), (" contiene un altro dizionario, quindi il primo accesso recupera il dizionario interno. Il secondo accesso, con la chiave ", False), ("\"nome\"", True), (", recupera il valore ", False), ("\"Anna\"", True), (".", False)])

    add_heading(doc, "Dati anagrafici con data nidificata", 2, prefix="II.")
    add_para(doc, "La data di nascita mostra perché la nidificazione è utile. Giorno, mese e anno appartengono allo stesso concetto, ma sono tre valori distinti. Inserirli dentro un dizionario interno permette di mantenere un'unica chiave esterna, data_nascita, senza perdere la possibilità di accedere ai singoli elementi.")
    add_code(doc, '''persona = {}
print("Inserire i seguenti dati anagrafici:")

persona["nome"] = input("Nome: ")
persona["cognome"] = input("Cognome: ")
persona["luogo_nascita"] = input("Luogo di nascita: ")

persona["data_nascita"] = {}
print("Data di nascita:")
persona["data_nascita"]["giorno"] = int(input("giorno: "))
persona["data_nascita"]["mese"] = int(input("mese: "))
persona["data_nascita"]["anno"] = int(input("anno: "))''')
    add_para(doc, [("Il dizionario ", False), ("persona", True), (" viene riempito prima con dati semplici, come nome, cognome e luogo di nascita. Poi la chiave ", False), ("\"data_nascita\"", True), (" riceve un dizionario vuoto, che diventa il contenitore dei tre campi della data. Le istruzioni successive inseriscono ", False), ("\"giorno\"", True), (", ", False), ("\"mese\"", True), (" e ", False), ("\"anno\"", True), (" nel dizionario interno, usando ", False), ("int(input(...))", True), (" per leggere valori numerici.", False)])

    add_heading(doc, "Liste di dizionari", 1, prefix="5.")
    add_para(doc, "Dopo avere costruito un dizionario per una persona, il passo successivo è gestire più persone. Invece di creare molte variabili separate, il materiale propone una lista in cui ogni elemento è un dizionario. La lista rappresenta l'insieme delle persone, mentre ciascun dizionario rappresenta una persona con la propria struttura interna.")
    add_para(doc, [("Il programma chiede prima quante persone devono essere inserite, poi usa un ciclo ", False), ("while", True), (" per ripetere la stessa procedura. A ogni iterazione viene creato un nuovo dizionario ", False), ("persona", True), (", riempito con i dati richiesti e aggiunto alla lista ", False), ("persone", True), (". La variabile ", False), ("n", True), (" conta quale persona si sta inserendo e permette al ciclo di fermarsi quando il numero richiesto è stato raggiunto.", False)])
    add_code(doc, '''persone = []
n_persone = int(input("Quante persone? "))
n = 1

while n <= n_persone:
    persona = {}
    print("Dati anagrafici della persona n. " + str(n) + ":")
    persona["nome"] = input(" nome: ")
    persona["cognome"] = input(" cognome: ")
    persona["luogo_nascita"] = input(" luogo di nascita: ")
    persona["data_nascita"] = {}
    print(" data di nascita:")
    persona["data_nascita"]["giorno"] = input("giorno: ")
    persona["data_nascita"]["mese"] = input(" mese: ")
    persona["data_nascita"]["anno"] = input(" anno: ")
    persone = persone + [persona]
    n += 1''')
    add_para(doc, [("La lista ", False), ("persone", True), (" nasce vuota e viene riempita gradualmente. Dentro il ciclo, ogni ", False), ("persona", True), (" è un dizionario autonomo, quindi i dati della prima persona non si mescolano con quelli della seconda. L'istruzione ", False), ("persone = persone + [persona]", True), (" aggiunge il dizionario appena costruito alla lista, mentre ", False), ("n += 1", True), (" prepara l'iterazione successiva.", False)])
    add_figure(doc, "slide_21_pic_1.png", "Esempio (6/4)", "Shell con inserimento di più persone in una lista di dizionari", 13.5)

    add_heading(doc, "Punti come dizionari in una lista", 2, prefix="I.")
    add_para(doc, "La stessa struttura viene applicata anche ai punti. Ogni punto ha due coordinate, x e y, quindi può essere rappresentato come un dizionario con due chiavi. Più punti vengono poi raccolti in una lista, esattamente come più persone vengono raccolte nella lista persone.")
    add_code(doc, '''n_punti = int(input("Numero di punti: "))
punti = []
n = 1

while n <= n_punti:
    punto = {}
    print("Coordinate del punto n. " + str(n) + ":")
    punto['x'] = int(input("x: "))
    punto['y'] = int(input("y: "))
    punti = punti + [punto]
    n = n + 1

print("Sono stati inseriti i seguenti punti:")
print(punti)''')
    add_para(doc, [("Il programma legge prima il numero di punti da acquisire e prepara una lista vuota chiamata ", False), ("punti", True), (". Nel ciclo viene creato un dizionario ", False), ("punto", True), (", quindi vengono lette le due coordinate e salvate nelle chiavi ", False), ("'x'", True), (" e ", False), ("'y'", True), (". Alla fine di ogni iterazione, il dizionario del punto viene aggiunto alla lista e il contatore avanza.", False)])

    add_heading(doc, "Liste come valori nei dizionari", 1, prefix="6.")
    add_para(doc, "Le liste non servono soltanto a contenere più dizionari; possono anche essere valori dentro un dizionario. Questa forma è utile quando un oggetto ha alcuni dati singoli e un gruppo di valori dello stesso tipo. Il caso dell'atleta nel salto in alto usa una lista per memorizzare le misure dei salti validi e non validi.")
    add_para(doc, [("Il dizionario dell'atleta contiene nome, cognome e numero di gara come valori singoli. La chiave ", False), ("\"salti\"", True), (" contiene invece una lista, perché il numero di misure può variare e perché tutte le misure appartengono allo stesso campo informativo. Il valore ", False), ("0", True), (" viene usato per indicare i salti non validi, secondo la convenzione presentata nel materiale.", False)])
    add_code(doc, '''atleta = {
    "nome": "Marco",
    "cognome": "Rossi",
    "numero_gara": 12,
    "salti": [185, 190, 0, 195, 0]
}

print("Atleta:", atleta["nome"], atleta["cognome"])
print("Numero gara:", atleta["numero_gara"])
print("Salti validi:", atleta["salti"])''')
    add_para(doc, [("Il dizionario viene creato direttamente con tutte le sue coppie chiave-valore. Le prime tre chiavi contengono informazioni semplici, mentre ", False), ("\"salti\"", True), (" contiene una lista di misure espresse in centimetri. Le istruzioni con ", False), ("print()", True), (" mostrano come accedere sia ai valori singoli sia alla lista completa associata alla chiave dei salti.", False)])

    add_heading(doc, "Parole chiave di un libro", 2, prefix="II.")
    add_para(doc, "Lo stesso schema compare nell'esempio della libreria. Un libro ha un titolo, un autore e un anno di pubblicazione, ma può avere più parole chiave. La lista consente di memorizzare questo insieme di parole dentro una sola chiave del dizionario.")
    add_code(doc, '''libro = {
    "titolo": "Python Base",
    "autore": "Mario Rossi",
    "anno": 2025,
    "parole_chiave": [
        "programmazione",
        "python",
        "informatica"
    ]
}

print("Titolo:", libro["titolo"])
print("Autore:", libro["autore"])
print("Anno:", libro["anno"])
print("Parole chiave:", libro["parole_chiave"])''')
    add_para(doc, [("Il dizionario ", False), ("libro", True), (" raccoglie dati di natura diversa ma collegati allo stesso oggetto. Le parole chiave sono inserite in una lista, perché sono più di una e possono essere lette insieme attraverso la chiave ", False), ("\"parole_chiave\"", True), (". Le istruzioni finali recuperano ogni campo con la stessa tecnica usata negli esempi precedenti, cioè l'accesso tramite chiave.", False)])

    add_heading(doc, "Coerenza delle strutture dati", 1, prefix="7.")
    add_para(doc, "Gli esempi presentati mostrano una progressione precisa. Prima si costruisce un dizionario semplice con dati inseriti dall'utente, poi si accede ai suoi valori, quindi si usano dizionari per rappresentare vettori, persone, punti, atleti e libri. La continuità tra gli esempi sta nell'idea che le chiavi descrivono il ruolo dei valori e rendono la struttura leggibile.")
    add_para(doc, [("Quando una struttura contiene altre strutture, l'accesso segue l'ordine in cui i dati sono organizzati. Per una lista dentro un dizionario si usa prima la chiave e poi l'indice; per un dizionario dentro un dizionario si usano prima la chiave esterna e poi quella interna. Per una lista di dizionari, invece, si seleziona un elemento della lista e poi si accede alle chiavi del dizionario selezionato.", False)])

    add_heading(doc, "Problemi tipici accompagnati da esempi", 2, prefix="I.")
    add_para(doc, [("Un primo problema riguarda la confusione tra stringhe e numeri quando si usa ", False), ("input()", True), (". L'esempio dell'età mostra che, se si vuole un numero intero, bisogna scrivere ", False), ("int(input(\"Età: \"))", True), (" invece di usare direttamente il valore prodotto da ", False), ("input()", True), (". La differenza si vede nel dizionario finale, perché l'età può comparire come testo oppure come numero.", False)])
    add_para(doc, [("Un secondo problema riguarda la scelta della struttura più adatta. L'esempio della data di nascita mostra che tre dati collegati possono essere raccolti in un dizionario interno, mentre l'esempio dei salti dell'atleta mostra che una serie di valori dello stesso tipo può essere raccolta in una lista. In entrambi i casi il programma resta leggibile perché ogni chiave conserva un significato chiaro.", False)])
    add_para(doc, [("Un terzo problema riguarda la ripetizione dei dati per più elementi. L'esempio delle persone e quello dei punti mostrano che si può creare un nuovo dizionario a ogni iterazione e inserirlo in una lista. In questo modo il programma non deve prevedere in anticipo variabili separate per ogni persona o per ogni punto.", False)])

    add_heading(doc, "Bibliografia", 1, numbered=False, prefix="")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run("Pensare da informatico Versione Python di Allen B. Downey, Jeffrey Elkner e Chris Meyers Traduzione di Alessandro Pocaterra, Capitoli 12, 13 e 14, Disponibili al link: ")
    set_run_font(r, "Times New Roman", 12, "000000", False)
    add_hyperlink(p, "https://www.python.it/doc/Howtothink/Howtothink-html-it/index.htm", "https://www.python.it/doc/Howtothink/Howtothink-html-it/index.htm")

    add_exercise_heading(doc)
    add_para(doc, "Esercizio 1. Scrivere un programma che chieda all'utente i dati di una persona e li memorizzi in un dizionario. Il dizionario deve contenere nome, cognome, luogo di nascita e data di nascita, dove la data deve essere organizzata come dizionario nidificato con giorno, mese e anno. L'esercizio richiede di usare le stesse idee viste negli esempi: input da tastiera, chiavi descrittive, conversione numerica dove necessaria e stampa finale della struttura.")
    add_para(doc, "Per svolgere l'esercizio, conviene costruire prima il dizionario esterno e verificare che nome, cognome e luogo siano salvati correttamente. Solo dopo si deve creare il dizionario interno della data di nascita e assegnarlo alla chiave appropriata. La consegna non richiede una soluzione già pronta, ma richiede di controllare che il risultato stampato mostri chiaramente la struttura nidificata.")
    add_para(doc, "Esercizio 2. Scrivere un programma che chieda quanti punti devono essere inseriti e memorizzi ogni punto come dizionario con chiavi x e y. Tutti i punti devono essere raccolti in una lista, seguendo lo schema in cui un ciclo crea un dizionario nuovo a ogni ripetizione. Alla fine il programma deve stampare la lista completa, così da verificare che il numero di dizionari corrisponda ai punti richiesti.")

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
